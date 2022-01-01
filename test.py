# This is a sample Python script.
from tkinter import *
from tkinter import ttk
# Press Shift+F10 to execute it or replace it with your code.
import sys
import time
import PIL.Image
import keyboard
import pyautogui
from datetime import date

from tkinter import *
from tkinter import ttk

import self as self
import tkcalendar
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from threading import Thread
import win32gui


class MeterMaidBot:
    shotFile = "C:/Users/jordan/Documents/Work/TempPics/shot.png"  # temporary image storage
    docxFile = "C:/Users/jordan/Documents/Work/PopScript.docx"  # main document
    hotkey = 'ctrl+shift+z'  # use this combination anytime while script is running

    def __init__(self, master):

        self.shotFile = "C:/Users/jordan/Documents/Work/TempPics/shot.png"  # temporary image storage
        self.docxFile = "C:/Users/jordan/Documents/Work/PopScript.docx"  # main document
        self.hotkey = 'ctrl+shift+z'  # use this combination anytime while script is running

        # Create a root window
        master.title('Meter Maid Bot')

        # Create a new child frame
        self.frame_content = ttk.Frame(master)
        self.frame_content.pack()

        # create labels
        ttk.Label(self.frame_content, text='Pops:').grid(row=1, column=0, columnspan=2, padx=5, sticky='sw')
        ttk.Label(self.frame_content, text='Dish #:').grid(row=3, column=0, columnspan=2, padx=5, sticky='sw')
        ttk.Label(self.frame_content, text='Choose date').grid(row=5, column=0, columnspan=2, padx=5, sticky='sw')

        # Create Pops Int Entry field
        self.spin1 = ttk.Spinbox(self.frame_content, style="My.TSpinbox", from_=1, to=48)

        # Create Dish String Entry field
        # Change to=attribute depending on how many dishes their are
        self.dishNum = ttk.Spinbox(self.frame_content, style="My.TSpinbox", from_=1, to=36)

        # Create Date entry
        self.v = StringVar(self.frame_content, tkcalendar.Calendar.date.today().strftime("%d/%m/%y"))
        self.cal1 = tkcalendar.Calendar(self.frame_content, selectmode = 'day',year = 2022, month = 1,
               day = 1, textvariable=self.v)

        # Format grid layout
        self.spin1.grid(row=2, column=0, padx=5, sticky='sw')
        self.dishNum.grid(row=4, column=0, padx=5, sticky='sw')
        self.cal1.grid(row=6, column=0, padx=5, sticky='sw')

        # Create buttons
        ttk.Button(self.frame_content, text="Submit", command=self.new_submit).grid(row=7, column=0, padx=5,
                                                                                    sticky='sw')

        # Style
        style = ttk.Style()
        style.theme_use('default')
        style.configure('My.TSpinbox', arrowsize=20)

        # self.model = "model"
        print("This is the constructor for test.py")
        # print(self)

    def new_submit(self):

        print("Recording saved" + " " + self.spin1.get())
        x = 0
        amount = []
        print("Pops: {}".format(self.spin1.get()))

        # Number of pops
        pops = self.spin1.get()
        pops = int(pops)

        # Name of Dish #
        dish = self.dishNum.get()

        print("Started. Waiting for", self.hotkey)

        while True:

            try:
                hwnd = win32gui.GetForegroundWindow()  # active window
                bbox = win32gui.GetWindowRect(hwnd)  # bounding rectangle

                # capture screen
                shot = pyautogui.screenshot(region=bbox)  # take screenshot, active app
                shot.save(self.shotFile)  # save screenshot

                for i in range(x, pops):

                    if keyboard.is_pressed(self.hotkey):

                        x += 1
                        amount.append(x)
                        print("you pressed the hotkey")
                        print(x)
                        print(amount)

                        if x % 2:

                            doc = Document(self.docxFile)  # open document

                            # Header - Ex. Dish 8

                            p = doc.add_paragraph()
                            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            r = p.add_run("Dish {}:".format(self.dishNum.get()))
                            r.font.size = Pt(24)
                            r.bold = False

                            p = doc.add_paragraph()
                            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            if x <= 2:
                                print("Cables 1 & 2")
                                r = p.add_run("Cables 1 & 2")
                                r.font.size = Pt(16)
                                r.bold = False

                            if (x > 2) & (x <= 4):
                                print("Cables 3 & 4")
                                r = p.add_run("Cables 3 & 4")
                                r.font.size = Pt(16)
                                r.bold = False

                            if (x > 4) & (x <= 6):
                                print("Cables 5 & 6")
                                r = p.add_run("Cables 5 & 6")
                                r.font.size = Pt(16)
                                r.bold = False

                            # Label 1 - Date:
                            doc.add_paragraph("Date: {}".format(self.v.get()))

                            print(str(self.v.get()))

                            # Label 2 - Location:

                            if x < pops / 4:
                                #doc.add_paragraph("Location: Outside")
                                print("less than a 1/4")

                            if x > pops / 4:
                                #doc.add_paragraph("Location: Inside")
                                print("more than 1 / 4")

                            doc.add_paragraph("Location: Outside")

                            # Label -3 - Status:


                            doc.add_paragraph("Status: Pre 5G-Filter")

                            # Sections - Horizontal
                            p = doc.add_paragraph()
                            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            r = p.add_run("Horizontal - ")
                            r.bold = True

                            doc.add_picture(self.shotFile, width=Inches(6.5))  # add image, default 6.5 inches wide

                            doc.save(self.docxFile)  # update document

                            print('Done capture.')
                            print("Odd number")

                        else:

                            doc = Document(self.docxFile)  # open document

                            # Sections - Horizontal
                            p = doc.add_paragraph()
                            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            r = p.add_run("Vertical - ")
                            r.bold = True

                            doc.add_picture(self.shotFile, width=Inches(6.5))  # add image, default 6.5 inches wide

                            doc.save(self.docxFile)  # update document
                            print("Even Number")

                        time.sleep(0.25)

            except Exception as e:  # allow program to keep running

                print("Capture Error:", e)
                return self

    def test_method(self):
        # mmb = MeterMaidBot()
        k = keyboard.get_hotkey_name(self.hotkey)
        print("Test Method ")
        print(k)
        return self

def main():

    root = Tk()
    app = MeterMaidBot(root)
    root.mainloop()



if __name__ == '__main__':
    main()
