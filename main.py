# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
import sys
import time
import PIL.Image
import keyboard
import pyautogui
from datetime import date

import self as self
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from threading import Thread
import win32gui


class MeterMaidBot:

    shotFile = "C:/Users/jordan/Documents/Work/TempPics/shot.png"  # temporary image storage
    docxFile = "C:/Users/jordan/Documents/Work/PopScript.docx"  # main document
    hotkey = 'ctrl+shift+z'  # use this combination anytime while script is running

    def __init__(self):
        self.model = "model"
        print("This is the constructor")
        print(self)

    def test_method(self):
        MMB = MeterMaidBot()
        k = keyboard.get_hotkey_name(MMB.hotkey)
        print("Test Method ")
        print(k)
        return self

    def picture(self):
        MMB = MeterMaidBot()
        x = 0
        amount = []
        pops = int(input("Enter the amount of POPS: "))

        while True:

            try:
                hwnd = win32gui.GetForegroundWindow()  # active window
                bbox = win32gui.GetWindowRect(hwnd)  # bounding rectangle

                # capture screen
                shot = pyautogui.screenshot(region=bbox)  # take screenshot, active app
                shot.save(MMB.shotFile)  # save screenshot

                # Edit Screenshot
                im = PIL.Image.open(MMB.shotFile)
                crop = im.crop((47, 138, 757, 693))
                crop.save(MMB.shotFile, quality=100)

                for i in range(x, pops):

                    if keyboard.is_pressed(MMB.hotkey):

                        x += 1
                        amount.append(x)
                        print("you pressed the hotkey")
                        print(x)
                        print(amount)

                        if x % 2:

                            doc = Document(MMB.docxFile)  # open document

                            # Header - Ex. Dish 8
                            p = doc.add_paragraph()
                            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            r = p.add_run("Dish 8")
                            r.font.size = Pt(24)
                            r.bold = False

                            # Sub-header - Ex. Cables 1 & 2
                            p = doc.add_paragraph()
                            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            r = p.add_run("Cables 1 & 2")
                            r.font.size = Pt(16)
                            r.bold = False

                            # Label 1 - Date:
                            doc.add_paragraph("Date: 12/9/2021")

                            # Label 2 - Location:
                            doc.add_paragraph("Location: Outside")

                            # Label -3 - Status:
                            doc.add_paragraph("Status: Pre 5G-Filter")

                            # Sections - Horizontal
                            p = doc.add_paragraph()
                            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            r = p.add_run("Horizontal - ")
                            r.bold = True

                            doc.add_picture(MMB.shotFile, width=Inches(6.5))  # add image, default 6.5 inches wide

                            doc.save(MMB.docxFile)  # update document

                            print('Done capture.')
                            print("Odd number")

                        else:

                            doc = Document(MMB.docxFile)  # open document

                            # Sections - Horizontal
                            p = doc.add_paragraph()
                            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            r = p.add_run("Vertical - ")
                            r.bold = True

                            doc.add_picture(MMB.shotFile, width=Inches(6.5))  # add image, default 6.5 inches wide

                            doc.save(MMB.docxFile)  # update document
                            print("Even Number")

                        time.sleep(0.25)

            except Exception as e:  # allow program to keep running

                print("Capture Error:", e)

    # def config():
    #     keyboard.add_hotkey(hotkey, MeterMaidBot.test_method)
    #     print("Started. Waiting for", hotkey)
    #     keyboard.wait()


def main():
    MeterMaidBot.picture(self)


if __name__ == '__main__':
    main()
